"""Microsoft Word DOCX document parser implementation."""

import logging
from typing import List, Dict, Any, Optional
import tempfile
import os
from pathlib import Path

from core.base import Parser, Document, ProcessingResult

logger = logging.getLogger(__name__)

try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    DocxDocument = None  # Type hint placeholder
    logger.debug("python-docx not available. DOCX parsing will not be available.")


class DocxParser(Parser):
    """Parser for Microsoft Word DOCX documents."""
    
    def __init__(self, name: str = "DocxParser", config: Optional[Dict[str, Any]] = None):
        super().__init__(name, config)
        config = config or {}
        self.extract_headers = config.get("extract_headers", True)
        self.extract_tables = config.get("extract_tables", True)
        self.extract_images = config.get("extract_images", False)
        self.preserve_structure = config.get("preserve_structure", True)
        self.include_metadata = config.get("include_metadata", True)
    
    def validate_config(self) -> bool:
        """Validate parser configuration."""
        if not HAS_PYTHON_DOCX:
            logger.error("python-docx package required for DOCX parsing")
            return False
        return True
    
    def parse(self, source: str, **kwargs) -> ProcessingResult:
        """Parse DOCX content into documents."""
        errors = []
        documents = []
        
        if not HAS_PYTHON_DOCX:
            logger.error("python-docx package required for DOCX parsing")
            errors.append({"error": "python-docx package not installed", "source": source})
            return ProcessingResult(documents=[], errors=errors)
        
        try:
            # Load DOCX document directly from file path
            doc = DocxDocument(source)
            documents = self._parse_docx_document(doc, source=source, **kwargs)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            errors.append({"error": str(e), "source": source})
            
        return ProcessingResult(documents=documents, errors=errors)
    
    def _parse_docx_document(self, doc: 'DocxDocument', **kwargs) -> List[Document]:
        """Parse the DOCX document object."""
        source = kwargs.get('source', 'unknown')
        
        if self.preserve_structure:
            return self._parse_by_sections(doc, source)
        else:
            return [self._create_single_document(doc, source)]
    
    def _parse_by_sections(self, doc: 'DocxDocument', source: str) -> List[Document]:
        """Parse DOCX by sections based on headings."""
        documents = []
        current_section = []
        current_header = ""
        section_count = 0
        
        for paragraph in doc.paragraphs:
            # Check if paragraph is a heading
            if paragraph.style.name.startswith('Heading'):
                # Save previous section if it has content
                if current_section:
                    section_text = '\n'.join(current_section).strip()
                    if section_text:
                        doc_obj = self._create_section_document(
                            section_text, current_header, section_count, source, doc
                        )
                        documents.append(doc_obj)
                        section_count += 1
                
                # Start new section
                current_header = paragraph.text.strip()
                current_section = [paragraph.text]
            else:
                # Add to current section
                text = paragraph.text.strip()
                if text:
                    current_section.append(text)
        
        # Handle last section
        if current_section:
            section_text = '\n'.join(current_section).strip()
            if section_text:
                doc_obj = self._create_section_document(
                    section_text, current_header, section_count, source, doc
                )
                documents.append(doc_obj)
        
        # If no sections found, create single document
        if not documents:
            documents = [self._create_single_document(doc, source)]
        
        return documents
    
    def _create_single_document(self, doc: 'DocxDocument', source: str) -> Document:
        """Create a single document from entire DOCX."""
        # Extract all text
        full_text = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
        
        content = '\n'.join(full_text)
        
        # Create metadata
        metadata = {
            "type": "docx_document",
            "source": source,
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
        
        # Extract document properties
        if self.include_metadata and hasattr(doc, 'core_properties'):
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
        
        # Extract structural metadata
        if self.extract_headers:
            headers = self._extract_headers(doc)
            if headers:
                metadata["headers"] = headers
                metadata["header_count"] = len(headers)
        
        if self.extract_tables:
            tables_info = self._extract_tables_info(doc)
            if tables_info:
                metadata["table_count"] = len(tables_info)
                metadata["tables"] = tables_info
        
        # Count paragraphs
        metadata["paragraph_count"] = len([p for p in doc.paragraphs if p.text.strip()])
        
        doc_id = f"docx_doc_{hash(content) % 10000}"
        
        return Document(
            id=doc_id,
            content=content,
            metadata=metadata,
            source=source
        )
    
    def _create_section_document(
        self, content: str, header: str, section_num: int, source: str, doc: 'DocxDocument'
    ) -> Document:
        """Create a document from a DOCX section."""
        metadata = {
            "type": "docx_section",
            "source": source,
            "section": section_num,
            "header": header,
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
        
        # Add document-level metadata to sections
        if self.include_metadata and hasattr(doc, 'core_properties'):
            core_props = doc.core_properties
            if core_props.title:
                metadata["document_title"] = core_props.title
            if core_props.author:
                metadata["document_author"] = core_props.author
        
        doc_id = f"docx_section_{section_num}_{hash(content) % 10000}"
        
        return Document(
            id=doc_id,
            content=content,
            metadata=metadata,
            source=source
        )
    
    def _extract_headers(self, doc: 'DocxDocument') -> List[Dict[str, Any]]:
        """Extract headers/headings from the document."""
        headers = []
        
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                # Extract heading level
                level_str = paragraph.style.name.replace('Heading ', '')
                try:
                    level = int(level_str)
                except ValueError:
                    level = 1
                
                headers.append({
                    "level": level,
                    "text": paragraph.text.strip(),
                    "style": paragraph.style.name
                })
        
        return headers
    
    def _extract_tables_info(self, doc: 'DocxDocument') -> List[Dict[str, Any]]:
        """Extract information about tables in the document."""
        tables_info = []
        
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0
            
            # Extract first few rows as sample data
            sample_data = []
            for row_idx, row in enumerate(table.rows[:3]):  # First 3 rows
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                sample_data.append(row_data)
            
            tables_info.append({
                "table_index": i,
                "rows": rows,
                "columns": cols,
                "sample_data": sample_data
            })
        
        return tables_info
    
    def _extract_images_info(self, doc: 'DocxDocument') -> List[Dict[str, Any]]:
        """Extract information about images in the document."""
        # This is more complex and requires access to the document relationships
        # For now, return basic count
        images_info = []
        
        # This would require deeper inspection of the DOCX structure
        # For basic implementation, we can count image elements if accessible
        try:
            # Access document part to get images
            document_part = doc.part
            if hasattr(document_part, 'related_parts'):
                image_parts = [
                    part for part in document_part.related_parts.values()
                    if 'image' in part.content_type
                ]
                
                for i, part in enumerate(image_parts):
                    images_info.append({
                        "image_index": i,
                        "content_type": part.content_type,
                        "size": len(part.blob) if hasattr(part, 'blob') else 0
                    })
        except Exception as e:
            logger.debug(f"Could not extract image info: {e}")
        
        return images_info
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """Get supported file extensions."""
        return ['.docx']
    
    @classmethod
    def get_description(cls) -> str:
        """Get parser description."""
        return "Microsoft Word DOCX document parser that extracts text, headers, tables, and metadata."